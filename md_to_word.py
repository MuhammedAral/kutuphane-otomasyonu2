# -*- coding: utf-8 -*-
"""
Markdown dosyasını Word'e dönüştürür (Birebir içerik)
"""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def md_to_docx(md_file, docx_file):
    """Markdown dosyasını Word'e dönüştürür"""
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Sayfa ayarları
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_content = []
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i]
        
        # Kod bloğu başlangıç/bitiş
        if line.strip().startswith('```'):
            if in_code_block:
                # Kod bloğu sonu - içeriği ekle
                p = doc.add_paragraph()
                code_text = '\n'.join(code_content)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Cm(0.5)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_content.append(line)
            i += 1
            continue
        
        # Tablo satırı kontrolü
        if line.strip().startswith('|') and line.strip().endswith('|'):
            # Ayırıcı satır mı kontrol et
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                i += 1
                continue
            
            # Tablo satırı
            cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            
            if not in_table:
                in_table = True
                table_rows = [cells]
            else:
                table_rows.append(cells)
            
            # Sonraki satır tablo mu kontrol et
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if not (next_line.startswith('|') and next_line.endswith('|')):
                    # Tablo bitti, oluştur
                    if table_rows:
                        num_cols = max(len(row) for row in table_rows)
                        table = doc.add_table(rows=len(table_rows), cols=num_cols)
                        table.style = 'Table Grid'
                        for row_idx, row_data in enumerate(table_rows):
                            for col_idx, cell_data in enumerate(row_data):
                                if col_idx < len(table.rows[row_idx].cells):
                                    table.rows[row_idx].cells[col_idx].text = cell_data
                                    if row_idx == 0:
                                        for para in table.rows[row_idx].cells[col_idx].paragraphs:
                                            for run in para.runs:
                                                run.bold = True
                        doc.add_paragraph()
                    table_rows = []
                    in_table = False
            i += 1
            continue
        
        # Tablo dışında normal satır, eğer tablo açıksa kapat
        if in_table:
            if table_rows:
                num_cols = max(len(row) for row in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.style = 'Table Grid'
                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < len(table.rows[row_idx].cells):
                            table.rows[row_idx].cells[col_idx].text = cell_data
                            if row_idx == 0:
                                for para in table.rows[row_idx].cells[col_idx].paragraphs:
                                    for run in para.runs:
                                        run.bold = True
                doc.add_paragraph()
            table_rows = []
            in_table = False
        
        # Başlıklar
        if line.startswith('# '):
            heading = doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        
        if line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue
        
        # Yatay çizgi
        if line.strip() == '---':
            doc.add_paragraph('─' * 50)
            i += 1
            continue
        
        # Madde işaretli liste
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue
        
        # Numaralı liste
        match = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if match:
            text = match.group(2)
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue
        
        # Boş satır
        if not line.strip():
            i += 1
            continue
        
        # Normal paragraf
        p = doc.add_paragraph()
        
        # Bold ve italic işleme
        text = line.strip()
        
        # **bold** ve *italic* işle
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run = p.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            else:
                p.add_run(part)
        
        i += 1
    
    # Dosyayı kaydet
    doc.save(docx_file)
    print(f"✅ Word dosyası oluşturuldu: {docx_file}")

if __name__ == "__main__":
    md_to_docx("KUTUPHANE_OTOMASYON_DOKUMANTASYONU.md", "KUTUPHANE_OTOMASYON_DOKUMANTASYONU.docx")
