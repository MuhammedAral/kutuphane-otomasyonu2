# -*- coding: utf-8 -*-
"""
Kütüphane Otomasyon Sistemi - Word Dokümantasyon Oluşturucu
Bu script, projein kapsamlı dokümantasyonunu Word (.docx) formatında oluşturur.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

def add_heading_1(doc, text):
    """Başlık 1 ekler"""
    heading = doc.add_heading(text, level=1)
    heading.style.font.color.rgb = RGBColor(0x00, 0x51, 0x8A)
    return heading

def add_heading_2(doc, text):
    """Başlık 2 ekler"""
    return doc.add_heading(text, level=2)

def add_heading_3(doc, text):
    """Başlık 3 ekler"""
    return doc.add_heading(text, level=3)

def add_table(doc, headers, rows):
    """Tablo ekler"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    # Data rows
    for row_data in rows:
        row = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row[i].text = str(cell_data)
    
    doc.add_paragraph()  # Boşluk
    return table

def add_code_block(doc, code, language=""):
    """Kod bloğu ekler"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p

def create_documentation():
    """Ana dokümantasyon oluşturma fonksiyonu"""
    
    doc = Document()
    
    # Sayfa ayarları
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4
    section.page_width = Cm(21)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    
    # ============ KAPAK SAYFASI ============
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    run = title.add_run("📚 KÜTÜPHANE OTOMASYON SİSTEMİ")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x00, 0x51, 0x8A)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    run = subtitle.add_run("Teknik Dokümantasyon")
    run.font.size = Pt(20)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Proje Adı: ").bold = True
    info.add_run("Kütüphane Otomasyon Sistemi\n\n")
    info.add_run("Geliştiriciler: ").bold = True
    info.add_run("Muhammed Ali Aral, Yağız Van\n\n")
    info.add_run("Tarih: ").bold = True
    info.add_run(f"{datetime.datetime.now().strftime('%d %B %Y')}\n\n")
    info.add_run("Versiyon: ").bold = True
    info.add_run("1.0")
    
    doc.add_page_break()
    
    # ============ İÇİNDEKİLER ============
    add_heading_1(doc, "İÇİNDEKİLER")
    
    toc = [
        "1. Proje Özeti",
        "2. Sistem Gereksinimleri",
        "3. Teknoloji Yığını",
        "4. Proje Yapısı",
        "5. Veritabanı Tasarımı",
        "6. REST API Dokümantasyonu",
        "7. WPF Masaüstü Uygulaması",
        "8. Güvenlik Özellikleri",
        "9. Kurulum ve Çalıştırma",
        "10. Kaynak Kod Detayları",
        "11. Web Sitesi (Web Arayüzü)"
    ]
    
    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(1)
    
    doc.add_page_break()
    
    # ============ 1. PROJE ÖZETİ ============
    add_heading_1(doc, "1. PROJE ÖZETİ")
    
    add_heading_2(doc, "1.1 Genel Bakış")
    doc.add_paragraph(
        "Kütüphane Otomasyon Sistemi, modern ve kullanıcı dostu bir kütüphane yönetim yazılımıdır. "
        "Sistem, WPF (Windows Presentation Foundation) teknolojisi ile geliştirilmiş masaüstü uygulaması "
        "ve ASP.NET Core ile geliştirilmiş REST API'den oluşmaktadır."
    )
    
    add_heading_2(doc, "1.2 Temel Özellikler")
    
    add_heading_3(doc, "Kullanıcı Sistemi")
    features = [
        "Yönetici Paneli: Tüm işlemlere tam erişim",
        "Üye Paneli: Kitap görüntüleme, değerlendirme ve kişisel ödünç takibi",
        "Güvenli giriş sistemi (SHA256 şifreleme + JWT)",
        "Gmail ile şifremi unuttum özelliği",
        "E-posta doğrulama ile kayıt"
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')
    
    add_heading_3(doc, "Kitap İşlemleri")
    features = [
        "Kitap ekleme, düzenleme ve silme",
        "Toplu kitap silme (akıllı ödünç kontrolü)",
        "Excel'den içe/dışa aktarma",
        "Barkod tarama ile hızlı işlem",
        "ISBN-10 ve ISBN-13 doğrulama",
        "Kitap türü yönetimi",
        "Stok takibi",
        "Kitap değerlendirme ve yorum sistemi"
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')
    
    add_heading_3(doc, "Üye İşlemleri")
    features = [
        "Yeni üye kaydı (Gmail doğrulama)",
        "Üye bilgilerini güncelleme",
        "Üyeleri aktif/pasif yapma",
        "Akıllı silme (ilişkili kayıtları temizler)"
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')
    
    add_heading_3(doc, "Ödünç İşlemleri")
    features = [
        "Kitap ödünç verme",
        "İade alma",
        "Geciken kitapları takip etme",
        "Gecikme ücreti hesaplama",
        "Filtreleme (Tümü, Ödünçte, Geciken, İade Edilmiş)"
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')
    
    add_heading_3(doc, "Raporlar ve İstatistikler")
    features = [
        "Dashboard istatistikleri (gerçek zamanlı)",
        "Geciken kitaplar listesi",
        "Excel rapor çıktısı"
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')
    
    doc.add_page_break()
    
    # ============ 2. SİSTEM GEREKSİNİMLERİ ============
    add_heading_1(doc, "2. SİSTEM GEREKSİNİMLERİ")
    
    add_heading_2(doc, "2.1 Geliştirme Ortamı")
    reqs = [".NET 8.0 SDK", "Visual Studio 2022 veya Visual Studio Code", "Git (versiyon kontrolü için)"]
    for r in reqs:
        doc.add_paragraph(r, style='List Bullet')
    
    add_heading_2(doc, "2.2 Çalıştırma Ortamı")
    reqs = [
        "Windows 10/11 (WPF uygulaması için)",
        "PostgreSQL veritabanı (Supabase üzerinden)",
        "İnternet bağlantısı (bulut veritabanı için)"
    ]
    for r in reqs:
        doc.add_paragraph(r, style='List Bullet')
    
    add_heading_2(doc, "2.3 API Bağımlılıkları (NuGet Paketleri)")
    add_table(doc, 
        ["Paket", "Versiyon", "Açıklama"],
        [
            ["MailKit", "4.14.1", "E-posta gönderimi"],
            ["Microsoft.AspNetCore.Authentication.JwtBearer", "8.0.0", "JWT kimlik doğrulama"],
            ["Microsoft.AspNetCore.OpenApi", "8.0.0", "OpenAPI desteği"],
            ["Swashbuckle.AspNetCore", "6.5.0", "Swagger UI"],
            ["Npgsql", "8.0.5", "PostgreSQL bağlantısı"]
        ]
    )
    
    add_heading_2(doc, "2.4 WPF Bağımlılıkları (NuGet Paketleri)")
    add_table(doc,
        ["Paket", "Versiyon", "Açıklama"],
        [
            ["MaterialDesignThemes", "4.9.0", "Material Design UI"],
            ["MaterialDesignColors", "2.1.4", "Material Design renk paleti"],
            ["ClosedXML", "0.102.2", "Excel işlemleri"],
            ["ZXing.Net", "0.16.11", "Barkod okuma"],
            ["ZXing.Net.Bindings.Windows.Compatibility", "0.16.14", "Windows barkod desteği"],
            ["AForge.Video.DirectShow", "2.2.5", "Kamera erişimi"],
            ["Npgsql", "8.0.5", "PostgreSQL bağlantısı"]
        ]
    )
    
    doc.add_page_break()
    
    # ============ 3. TEKNOLOJİ YIĞINI ============
    add_heading_1(doc, "3. TEKNOLOJİ YIĞINI")
    
    add_table(doc,
        ["Bileşen", "Teknoloji"],
        [
            ["Masaüstü Uygulama", ".NET 8.0 WPF + Material Design"],
            ["REST API", "ASP.NET Core 8.0 Minimal API"],
            ["Veritabanı", "PostgreSQL (Supabase)"],
            ["Authentication", "JWT Bearer Token"],
            ["Excel İşlemleri", "ClosedXML"],
            ["Barkod", "ZXing.Net + AForge.Video"],
            ["E-posta", "MailKit (Gmail SMTP)"],
            ["Rate Limiting", "System.Threading.RateLimiting"]
        ]
    )
    
    doc.add_page_break()
    
    # ============ 4. PROJE YAPISI ============
    add_heading_1(doc, "4. PROJE YAPISI")
    
    project_structure = """kutuphane-otomasyonu/
├── api/                          # REST API projesi
│   ├── Program.cs                # API endpoint'leri (1900+ satır)
│   ├── Services/                 # Servis katmanı
│   │   ├── EmailService.cs       # E-posta gönderim servisi
│   │   └── IEmailService.cs      # E-posta servis arayüzü
│   ├── appsettings.json          # JWT ve Email yapılandırması
│   ├── KutuphaneApi.csproj       # API proje dosyası
│   └── Dockerfile                # Docker yapılandırması
│
├── csharp/                       # WPF masaüstü uygulaması
│   ├── Views/                    # Ana pencereler
│   │   ├── LoginWindow.xaml(.cs) # Giriş ekranı
│   │   ├── RegisterWindow.xaml(.cs) # Kayıt ekranı
│   │   ├── AdminWindow.xaml(.cs) # Yönetici paneli
│   │   ├── MemberWindow.xaml(.cs) # Üye paneli
│   │   └── ForgotPasswordWindow.xaml(.cs) # Şifre sıfırlama
│   │
│   ├── Pages/                    # Yönetici sayfaları
│   │   ├── DashboardPage.xaml(.cs) # Ana panel
│   │   ├── KitaplarPage.xaml(.cs) # Kitap yönetimi
│   │   ├── KitapDialog.xaml(.cs) # Kitap ekleme/düzenleme
│   │   ├── KitapDetayDialog.xaml(.cs) # Kitap detay ve yorumlar
│   │   ├── UyelerPage.xaml(.cs) # Üye yönetimi
│   │   ├── UyeDialog.xaml(.cs) # Üye ekleme
│   │   ├── OduncPage.xaml(.cs) # Ödünç işlemleri
│   │   ├── OduncDialog.xaml(.cs) # Ödünç verme
│   │   ├── RaporlarPage.xaml(.cs) # Raporlar
│   │   ├── AyarlarPage.xaml(.cs) # Sistem ayarları
│   │   └── BarcodeScannerDialog.xaml(.cs) # Barkod tarama
│   │
│   ├── MemberPages/              # Üye sayfaları
│   │   ├── AnasayfaPage.xaml(.cs) # Üye ana sayfa
│   │   ├── KitaplarViewPage.xaml(.cs) # Kitap listesi
│   │   ├── OdunclerimPage.xaml(.cs) # Ödünçlerim
│   │   └── ProfilPage.xaml(.cs) # Profil bilgileri
│   │
│   ├── Assets/                   # Görseller
│   ├── ApiService.cs             # API iletişim servisi
│   ├── DatabaseHelper.cs         # Veritabanı yardımcı sınıfı
│   ├── CurrentSession.cs         # Oturum bilgileri
│   └── KutuphaneOtomasyon.csproj # WPF proje dosyası
│
├── docker-compose.yml            # Docker Compose yapılandırması
├── kutuphane-otomasyonu.sln      # Solution dosyası
└── README.md                     # Proje açıklaması"""
    
    add_code_block(doc, project_structure)
    
    doc.add_page_break()
    
    # ============ 5. VERİTABANI TASARIMI ============
    add_heading_1(doc, "5. VERİTABANI TASARIMI")
    
    add_heading_2(doc, "5.1 Veritabanı Bağlantısı")
    doc.add_paragraph(
        "Sistem, Supabase üzerinde barındırılan PostgreSQL veritabanını kullanmaktadır."
    )
    
    add_table(doc,
        ["Parametre", "Değer"],
        [
            ["Host", "aws-1-eu-central-1.pooler.supabase.com"],
            ["Port", "6543"],
            ["Database", "postgres"],
            ["SSL Mode", "Require"]
        ]
    )
    
    add_heading_2(doc, "5.2 Tablo Yapıları")
    
    add_heading_3(doc, "5.2.1 Kullanicilar Tablosu")
    kullanicilar_sql = """CREATE TABLE Kullanicilar (
    KullaniciID SERIAL PRIMARY KEY,
    KullaniciAdi VARCHAR(50) UNIQUE NOT NULL,
    Sifre VARCHAR(256) NOT NULL,
    AdSoyad VARCHAR(100) NOT NULL,
    Email VARCHAR(100),
    Telefon VARCHAR(20),
    Rol VARCHAR(20) DEFAULT 'Uye',
    AktifMi BOOLEAN DEFAULT TRUE,
    OlusturmaTarihi TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);"""
    add_code_block(doc, kullanicilar_sql)
    
    add_table(doc,
        ["Alan", "Tip", "Açıklama"],
        [
            ["KullaniciID", "SERIAL", "Birincil anahtar, otomatik artan"],
            ["KullaniciAdi", "VARCHAR(50)", "Benzersiz kullanıcı adı"],
            ["Sifre", "VARCHAR(256)", "SHA256 ile hashlenmiş şifre"],
            ["AdSoyad", "VARCHAR(100)", "Kullanıcının tam adı"],
            ["Email", "VARCHAR(100)", "E-posta adresi"],
            ["Telefon", "VARCHAR(20)", "Telefon numarası"],
            ["Rol", "VARCHAR(20)", "'Yonetici' veya 'Uye'"],
            ["AktifMi", "BOOLEAN", "Hesap aktif mi?"],
            ["OlusturmaTarihi", "TIMESTAMP", "Kayıt tarihi"]
        ]
    )
    
    add_heading_3(doc, "5.2.2 Kitaplar Tablosu")
    kitaplar_sql = """CREATE TABLE Kitaplar (
    KitapID SERIAL PRIMARY KEY,
    Baslik VARCHAR(200) NOT NULL,
    Yazar VARCHAR(100) NOT NULL,
    ISBN VARCHAR(20),
    Barkod VARCHAR(50),
    YayinYili INTEGER,
    TurID INTEGER REFERENCES KitapTurleri(TurID),
    StokAdedi INTEGER DEFAULT 1,
    MevcutAdet INTEGER DEFAULT 1,
    RafNo VARCHAR(20),
    SiraNo VARCHAR(20),
    Aciklama VARCHAR(500),
    EklenmeTarihi TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);"""
    add_code_block(doc, kitaplar_sql)
    
    add_heading_3(doc, "5.2.3 OduncIslemleri Tablosu")
    odunc_sql = """CREATE TABLE OduncIslemleri (
    IslemID SERIAL PRIMARY KEY,
    KitapID INTEGER REFERENCES Kitaplar(KitapID),
    UyeID INTEGER REFERENCES Kullanicilar(KullaniciID),
    OduncTarihi TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    BeklenenIadeTarihi TIMESTAMP,
    IadeTarihi TIMESTAMP,
    Durum VARCHAR(20) DEFAULT 'Odunc',
    CezaMiktari DECIMAL(10,2) DEFAULT 0
);"""
    add_code_block(doc, odunc_sql)
    
    add_heading_3(doc, "5.2.4 Diğer Tablolar")
    doc.add_paragraph("KitapTurleri: Kitap türlerini saklar (Roman, Hikaye, Şiir, vb.)", style='List Bullet')
    doc.add_paragraph("Degerlendirmeler: Kitap puanlama ve yorumları saklar", style='List Bullet')
    doc.add_paragraph("SifreSifirlamaIslemleri: Şifre sıfırlama kodlarını saklar", style='List Bullet')
    doc.add_paragraph("Ayarlar: Sistem ayarlarını saklar (gecikme ücreti, ödünç süresi)", style='List Bullet')
    
    doc.add_page_break()
    
    # ============ 6. REST API DOKÜMANTASYONU ============
    add_heading_1(doc, "6. REST API DOKÜMANTASYONU")
    
    add_heading_2(doc, "6.1 Genel Bilgiler")
    add_table(doc,
        ["Parametre", "Değer"],
        [
            ["Base URL", "http://localhost:5026"],
            ["Swagger UI", "http://localhost:5026/swagger"],
            ["Authentication", "JWT Bearer Token"]
        ]
    )
    
    add_heading_2(doc, "6.2 API Endpoint'leri")
    
    add_heading_3(doc, "Kimlik Doğrulama")
    add_table(doc,
        ["Method", "Endpoint", "Açıklama"],
        [
            ["POST", "/api/giris", "Giriş yap"],
            ["POST", "/api/auth/register", "Yeni kayıt"],
            ["POST", "/api/auth/verify-email", "E-posta doğrulama"],
            ["POST", "/api/auth/sifremi-unuttum", "Şifre sıfırlama kodu gönder"],
            ["POST", "/api/auth/sifre-sifirla", "Yeni şifre belirle"]
        ]
    )
    
    add_heading_3(doc, "Dashboard")
    add_table(doc,
        ["Method", "Endpoint", "Açıklama"],
        [
            ["GET", "/api/dashboard/stats", "İstatistikler"],
            ["GET", "/api/dashboard/geciken-kitaplar", "Geciken kitaplar listesi"]
        ]
    )
    
    add_heading_3(doc, "Kitaplar")
    add_table(doc,
        ["Method", "Endpoint", "Açıklama"],
        [
            ["GET", "/api/kitaplar", "Kitap listesi"],
            ["GET", "/api/kitaplar/{id}", "Kitap detayı"],
            ["POST", "/api/kitaplar", "Kitap ekle"],
            ["PUT", "/api/kitaplar/{id}", "Kitap güncelle"],
            ["DELETE", "/api/kitaplar/{id}", "Kitap sil"],
            ["DELETE", "/api/kitaplar/toplu", "Toplu kitap silme"],
            ["POST", "/api/kitaplar/toplu", "Toplu kitap ekleme"]
        ]
    )
    
    add_heading_3(doc, "Üyeler")
    add_table(doc,
        ["Method", "Endpoint", "Açıklama"],
        [
            ["GET", "/api/uyeler", "Üye listesi"],
            ["GET", "/api/uyeler/{id}", "Üye detayı"],
            ["POST", "/api/uyeler", "Üye ekle"],
            ["DELETE", "/api/uyeler/{id}", "Üye sil"]
        ]
    )
    
    add_heading_3(doc, "Ödünç İşlemleri")
    add_table(doc,
        ["Method", "Endpoint", "Açıklama"],
        [
            ["GET", "/api/odunc", "Ödünç listesi"],
            ["POST", "/api/odunc", "Ödünç ver"],
            ["PUT", "/api/odunc/{id}/iade", "İade al"],
            ["GET", "/api/odunc/stats", "Ödünç istatistikleri"],
            ["GET", "/api/odunc/uye/{uyeId}", "Üye ödünçleri"]
        ]
    )
    
    add_heading_3(doc, "Değerlendirmeler")
    add_table(doc,
        ["Method", "Endpoint", "Açıklama"],
        [
            ["GET", "/api/kitaplar/{id}/degerlendirmeler", "Kitap değerlendirmeleri"],
            ["GET", "/api/kitaplar/{id}/puan", "Ortalama puan"],
            ["DELETE", "/api/degerlendirmeler/{id}", "Değerlendirme sil"]
        ]
    )
    
    add_heading_3(doc, "Raporlar ve Diğer")
    add_table(doc,
        ["Method", "Endpoint", "Açıklama"],
        [
            ["GET", "/api/raporlar", "Detaylı raporlar"],
            ["GET", "/api/turler", "Kitap türleri"],
            ["GET", "/api/istatistikler", "Genel istatistikler"]
        ]
    )
    
    doc.add_page_break()
    
    # ============ 7. WPF MASAÜSTÜ UYGULAMASI ============
    add_heading_1(doc, "7. WPF MASAÜSTÜ UYGULAMASI")
    
    add_heading_2(doc, "7.1 Uygulama Mimarisi")
    doc.add_paragraph(
        "Uygulama, MVVM benzeri bir yapı kullanmaktadır:"
    )
    doc.add_paragraph("Views: Ana pencereler (Login, Register, Admin, Member)", style='List Bullet')
    doc.add_paragraph("Pages: Sayfa içerikleri (Dashboard, Kitaplar, Üyeler, vb.)", style='List Bullet')
    doc.add_paragraph("Services: API iletişim katmanı (ApiService)", style='List Bullet')
    doc.add_paragraph("Helpers: Yardımcı sınıflar (DatabaseHelper, CurrentSession)", style='List Bullet')
    
    add_heading_2(doc, "7.2 Views (Ana Pencereler)")
    
    add_table(doc,
        ["Dosya", "Açıklama"],
        [
            ["LoginWindow", "Kullanıcı giriş ekranı"],
            ["RegisterWindow", "Yeni kullanıcı kaydı ve e-posta doğrulama"],
            ["AdminWindow", "Yönetici paneli (sol menü + içerik alanı)"],
            ["MemberWindow", "Üye paneli (kısıtlı erişim)"],
            ["ForgotPasswordWindow", "Şifre sıfırlama ekranı"]
        ]
    )
    
    add_heading_2(doc, "7.3 Pages (Yönetici Sayfaları)")
    
    add_table(doc,
        ["Dosya", "Açıklama"],
        [
            ["DashboardPage", "Ana panel - istatistikler ve hızlı işlemler"],
            ["KitaplarPage", "Kitap yönetimi - CRUD, Excel import/export, toplu silme"],
            ["KitapDialog", "Kitap ekleme/düzenleme formu"],
            ["KitapDetayDialog", "Kitap detayları ve yorum sistemi"],
            ["UyelerPage", "Üye yönetimi - listeleme ve silme"],
            ["UyeDialog", "Yeni üye ekleme formu"],
            ["OduncPage", "Ödünç işlemleri - listeleme, iade alma"],
            ["OduncDialog", "Yeni ödünç verme formu"],
            ["RaporlarPage", "Detaylı raporlar ve istatistikler"],
            ["AyarlarPage", "Sistem ayarları"],
            ["BarcodeScannerDialog", "Kamera ile barkod tarama"]
        ]
    )
    
    add_heading_2(doc, "7.4 MemberPages (Üye Sayfaları)")
    
    add_table(doc,
        ["Dosya", "Açıklama"],
        [
            ["AnasayfaPage", "Üye ana sayfası - kişisel istatistikler"],
            ["KitaplarViewPage", "Kitap listesi görüntüleme"],
            ["OdunclerimPage", "Kullanıcının ödünç aldığı kitaplar"],
            ["ProfilPage", "Profil bilgileri ve güncelleme"]
        ]
    )
    
    add_heading_2(doc, "7.5 Servis Sınıfları")
    
    add_heading_3(doc, "ApiService.cs (650 satır)")
    doc.add_paragraph(
        "Tüm API çağrılarını yöneten statik sınıf. HttpClient yönetimi, JWT token yönetimi, "
        "JSON serialization ve tüm endpoint'ler için async metotlar içerir."
    )
    
    add_heading_3(doc, "DatabaseHelper.cs")
    doc.add_paragraph(
        "Veritabanı işlemleri için yardımcı sınıf. Bağlantı yönetimi, şifre hashleme (SHA256), "
        "ayar okuma/yazma işlemlerini gerçekleştirir."
    )
    
    add_heading_3(doc, "CurrentSession.cs")
    doc.add_paragraph(
        "Aktif kullanıcı oturum bilgilerini (UserId, AdSoyad, Rol) tutan statik sınıf."
    )
    
    doc.add_page_break()
    
    # ============ 8. GÜVENLİK ÖZELLİKLERİ ============
    add_heading_1(doc, "8. GÜVENLİK ÖZELLİKLERİ")
    
    add_heading_2(doc, "8.1 Şifre Güvenliği")
    doc.add_paragraph("SHA256 algoritması ile hashleme", style='List Bullet')
    doc.add_paragraph("Veritabanında düz metin şifre saklanmaz", style='List Bullet')
    doc.add_paragraph("Minimum 6 karakter zorunluluğu", style='List Bullet')
    
    add_heading_2(doc, "8.2 JWT Authentication")
    doc.add_paragraph("2 saat geçerlilik süresi", style='List Bullet')
    doc.add_paragraph("HMAC-SHA256 imzalama", style='List Bullet')
    doc.add_paragraph("Issuer ve Audience doğrulama", style='List Bullet')
    doc.add_paragraph("Token içeriği: NameIdentifier, Name, Role claims", style='List Bullet')
    
    add_heading_2(doc, "8.3 E-posta Doğrulama")
    doc.add_paragraph("Sadece @gmail.com kabul edilir", style='List Bullet')
    doc.add_paragraph("6 haneli rastgele doğrulama kodu", style='List Bullet')
    doc.add_paragraph("15 dakika geçerlilik süresi", style='List Bullet')
    
    add_heading_2(doc, "8.4 ISBN Doğrulama")
    doc.add_paragraph(
        "Hem ISBN-10 hem ISBN-13 formatları desteklenir ve check digit algoritması ile doğrulanır."
    )
    
    add_heading_2(doc, "8.5 SQL Injection Koruması")
    doc.add_paragraph(
        "Tüm veritabanı sorguları parametreli olarak yazılmıştır."
    )
    
    add_heading_2(doc, "8.6 Yetkilendirme Matrisi")
    add_table(doc,
        ["İşlem", "Yönetici", "Üye"],
        [
            ["Kitap Ekleme/Düzenleme/Silme", "✅", "❌"],
            ["Üye Ekleme/Silme", "✅", "❌"],
            ["Ödünç Verme/İade Alma", "✅", "❌"],
            ["Kitap Görüntüleme", "✅", "✅"],
            ["Yorum Yapma", "✅", "✅"],
            ["Kendi Yorumunu Silme", "✅", "✅"],
            ["Başkasının Yorumunu Silme", "✅", "❌"],
            ["Profil Güncelleme", "✅", "✅ (Kendi)"]
        ]
    )
    
    add_heading_2(doc, "8.7 API Rate Limiting (İstek Sınırlama)")
    doc.add_paragraph(
        "API güvenliğini ve stabilitesini korumak için rate limiting (hız sınırlama) uygulanmıştır."
    )
    
    add_table(doc,
        ["Özellik", "Değer"],
        [
            ["Algoritma", "Fixed Window (Sabit Pencere)"],
            ["Limit", "Her IP adresi için dakikada 100 istek"],
            ["Sıra (Queue)", "0 (Sıraya alınmaz, direkt reddedilir)"],
            ["Pencere", "1 Dakika"]
        ]
    )
    
    doc.add_paragraph("Limit aşıldığında:", style='List Bullet')
    doc.add_paragraph("  - HTTP Kodu: 429 Too Many Requests")
    doc.add_paragraph("  - Header: Retry-After: 60")
    doc.add_paragraph("  - Mesaj: 'Çok fazla istek gönderdiniz. Lütfen 1 dakika bekleyip tekrar deneyin.'")
    
    add_heading_3(doc, "Implementasyon Kodu")
    rate_limit_code = """builder.Services.AddRateLimiter(options =>
{
    options.GlobalLimiter = PartitionedRateLimiter.Create<HttpContext, string>(context =>
    {
        return RateLimitPartition.GetFixedWindowLimiter(
            partitionKey: context.Connection.RemoteIpAddress?.ToString() ?? "unknown",
            factory: _ => new FixedWindowRateLimiterOptions
            {
                AutoReplenishment = true,
                PermitLimit = 100,
                QueueLimit = 0,
                Window = TimeSpan.FromMinutes(1)
            });
    });
});"""
    add_code_block(doc, rate_limit_code)
    
    doc.add_page_break()
    
    # ============ 9. KURULUM VE ÇALIŞTIRMA ============
    add_heading_1(doc, "9. KURULUM VE ÇALIŞTIRMA")
    
    add_heading_2(doc, "9.1 Gereksinimler")
    doc.add_paragraph(".NET 8.0 SDK", style='List Bullet')
    doc.add_paragraph("İnternet bağlantısı (Supabase veritabanı için)", style='List Bullet')
    doc.add_paragraph("Windows 10/11 (WPF uygulaması için)", style='List Bullet')
    
    add_heading_2(doc, "9.2 API'yi Başlatma")
    add_code_block(doc, "cd api\ndotnet run")
    
    add_table(doc,
        ["Adres", "Açıklama"],
        [
            ["http://localhost:5026", "API Base URL"],
            ["http://localhost:5026/swagger", "Swagger UI"]
        ]
    )
    
    add_heading_2(doc, "9.3 WPF Uygulamasını Başlatma")
    add_code_block(doc, "cd csharp\ndotnet run")
    
    add_heading_2(doc, "9.4 Varsayılan Giriş Bilgileri")
    add_table(doc,
        ["Alan", "Değer"],
        [
            ["Kullanıcı Adı", "admin"],
            ["Şifre", "admin123"]
        ]
    )
    
    add_heading_2(doc, "9.5 EXE Oluşturma (Tek Dosya)")
    add_code_block(doc, 
        "cd csharp\ndotnet publish -c Release -r win-x64 --self-contained true -p:PublishSingleFile=true -o ./publish"
    )
    
    doc.add_page_break()
    
    # ============ 10. KAYNAK KOD DETAYLARI ============
    add_heading_1(doc, "10. KAYNAK KOD DETAYLARI")
    
    add_heading_2(doc, "10.1 API Program.cs Özeti (1906 satır)")
    doc.add_paragraph(
        "Program.cs dosyası, ASP.NET Core Minimal API yapısını kullanmaktadır."
    )
    
    doc.add_paragraph("Servis Yapılandırması (Satır 1-75): Swagger, CORS, JWT Authentication", style='List Bullet')
    doc.add_paragraph("Veritabanı İlklendirme (170-335): Tablo oluşturma, varsayılan veriler", style='List Bullet')
    doc.add_paragraph("Giriş API (390-444): Login işlemi", style='List Bullet')
    doc.add_paragraph("Dashboard API (446-517): İstatistikler ve geciken kitaplar", style='List Bullet')
    doc.add_paragraph("Şifre Sıfırlama API (519-703): Şifre işlemleri", style='List Bullet')
    doc.add_paragraph("Kitaplar API (705-932): CRUD işlemleri", style='List Bullet')
    doc.add_paragraph("Üyeler API (934-1532): CRUD işlemleri", style='List Bullet')
    doc.add_paragraph("Ödünç İşlemleri API (964-1220): Ödünç ve iade", style='List Bullet')
    doc.add_paragraph("Değerlendirmeler API (1242-1371): Puan ve yorumlar", style='List Bullet')
    doc.add_paragraph("Raporlar API (1767-1836): Detaylı raporlar", style='List Bullet')
    doc.add_paragraph("Request Models (1893-1906): DTO sınıfları", style='List Bullet')
    
    add_heading_2(doc, "10.2 Request/Response Modelleri")
    models_code = """public record LoginRequest(string KullaniciAdi, string Sifre);
public record KitapRequest(string Baslik, string Yazar, string? ISBN, 
    int? YayinYili, int? TurID, int? StokAdedi, string? RafNo);
public record OduncRequest(int KitapID, int UyeID, int? OduncGunu);
public record RegisterRequest(string KullaniciAdi, string Sifre, 
    string AdSoyad, string Email, string? Telefon);
public record UyeRequest(string KullaniciAdi, string Sifre, 
    string AdSoyad, string? Email, string? Telefon);
public record DegerlendirmeRequest(int KitapID, int UyeID, 
    int Puan, string? Yorum);"""
    add_code_block(doc, models_code)
    
    add_heading_2(doc, "10.3 HTTP Durum Kodları")
    add_table(doc,
        ["HTTP Kodu", "Anlamı"],
        [
            ["200", "Başarılı"],
            ["201", "Oluşturuldu"],
            ["400", "Geçersiz istek"],
            ["401", "Yetkisiz erişim"],
            ["403", "Yasaklı"],
            ["404", "Bulunamadı"],
            ["500", "Sunucu hatası"]
        ]
    )
    
    doc.add_page_break()
    
    # ============ 11. WEB SİTESİ ============
    add_heading_1(doc, "11. WEB SİTESİ (WEB ARAYÜZÜ)")
    
    add_heading_2(doc, "11.1 Genel Bakış")
    doc.add_paragraph(
        "Proje, masaüstü uygulamasının yanı sıra modern ve responsive bir web arayüzü de içermektedir. "
        "Web sitesi, API servisi tarafından sunulur ve tarayıcı üzerinden erişilebilir."
    )
    doc.add_paragraph("Erişim: http://localhost:5026 (API çalıştırıldığında)")
    
    add_heading_2(doc, "11.2 Web Sitesi Yapısı")
    web_structure = """website/
├── css/
│   └── styles.css              # Ana stil dosyası (1100+ satır)
├── js/
│   └── api.js                  # API iletişim ve yardımcı fonksiyonlar
├── admin/                      # Yönetici paneli
│   ├── index.html              # Admin ana sayfa
│   ├── kitaplar.html           # Kitap yönetimi
│   ├── uyeler.html             # Üye yönetimi
│   └── odunc.html              # Ödünç işlemleri
├── index.html                  # Üye ana sayfa
├── login.html                  # Giriş sayfası
├── kitaplar.html               # Kitap listesi (üye görünümü)
├── odunclerim.html             # Kullanıcının ödünçleri
└── profil.html                 # Profil sayfası"""
    add_code_block(doc, web_structure)
    
    add_heading_2(doc, "11.3 Tasarım Özellikleri")
    
    add_heading_3(doc, "Renk Paleti (Masaüstü ile Uyumlu)")
    add_table(doc,
        ["Değişken", "Renk Kodu", "Kullanım"],
        [
            ["--bg-primary", "#1e1e2e", "Ana arka plan"],
            ["--bg-card", "#2d2d44", "Kart arka planı"],
            ["--gradient-start", "#3b82f6", "Gradient başlangıcı (mavi)"],
            ["--gradient-end", "#8b5cf6", "Gradient bitişi (mor)"],
            ["--primary", "#3b82f6", "Ana renk"],
            ["--success", "#10b981", "Başarı durumu"],
            ["--warning", "#f59e0b", "Uyarı durumu"],
            ["--danger", "#ef4444", "Hata durumu"]
        ]
    )
    
    add_heading_2(doc, "11.4 Sayfa Detayları")
    
    add_heading_3(doc, "login.html - Giriş Sayfası")
    doc.add_paragraph("Kullanıcı adı ve şifre girişi", style='List Bullet')
    doc.add_paragraph("Şifremi Unuttum modal penceresi", style='List Bullet')
    doc.add_paragraph("Kayıt Ol modal penceresi", style='List Bullet')
    doc.add_paragraph("E-posta doğrulama modal penceresi", style='List Bullet')
    doc.add_paragraph("Rol bazlı yönlendirme (Admin → /admin/, Üye → /)", style='List Bullet')
    
    add_heading_3(doc, "index.html - Üye Ana Sayfa")
    doc.add_paragraph("İstatistik kartları (Toplam Kitap, Ödünçteki, Geciken)", style='List Bullet')
    doc.add_paragraph("Animasyonlu sayaçlar", style='List Bullet')
    doc.add_paragraph("Ödünçteki kitaplar tablosu", style='List Bullet')
    
    add_heading_3(doc, "kitaplar.html - Kitap Listesi")
    doc.add_paragraph("Kitap kartları grid görünümü", style='List Bullet')
    doc.add_paragraph("Arama fonksiyonu", style='List Bullet')
    doc.add_paragraph("Kitap detay modal'ı", style='List Bullet')
    doc.add_paragraph("Puan ve yorum sistemi", style='List Bullet')
    
    add_heading_2(doc, "11.5 Admin Paneli (/admin/)")
    add_table(doc,
        ["Sayfa", "Açıklama"],
        [
            ["admin/index.html", "Yönetici ana sayfa - istatistikler ve son işlemler"],
            ["admin/kitaplar.html", "Kitap yönetimi - CRUD işlemleri"],
            ["admin/uyeler.html", "Üye yönetimi"],
            ["admin/odunc.html", "Ödünç işlemleri - filtre ve iade"]
        ]
    )
    
    add_heading_2(doc, "11.6 JavaScript API Modülü (api.js)")
    
    doc.add_paragraph("Auth Nesnesi:", style='List Bullet')
    doc.add_paragraph("  - getToken(), setToken(), removeToken()")
    doc.add_paragraph("  - getUser(), setUser(), isLoggedIn()")
    doc.add_paragraph("  - logout(), requireAuth(), requireRole()")
    
    doc.add_paragraph()
    doc.add_paragraph("API Nesnesi:", style='List Bullet')
    doc.add_paragraph("  - login(), getKitaplar(), getKitap()")
    doc.add_paragraph("  - getTurler(), getOdunclerim(), getAllOdunc()")
    doc.add_paragraph("  - getProfilBilgileri(), getUyeler(), getIstatistikler()")
    doc.add_paragraph("  - degerlendirmeEkle(), degerlendirmeSil()")
    
    add_heading_2(doc, "11.7 CSS Stilleri (styles.css)")
    
    add_table(doc,
        ["Sınıf", "Görünüm"],
        [
            [".btn-gradient", "Mavi-mor gradient, gölgeli"],
            [".btn-outline", "Şeffaf arka plan, kenarlıklı"],
            [".stat-card", "İstatistik kartı (ikon + değer + etiket)"],
            [".table-card", "Tablo kartı (başlık + içerik)"],
            [".badge-success/warning/danger", "Durum göstergeleri"]
        ]
    )
    
    add_heading_2(doc, "11.8 Responsive Tasarım")
    doc.add_paragraph("768px altında sidebar gizlenir", style='List Bullet')
    doc.add_paragraph("480px altında grid tek sütuna düşer", style='List Bullet')
    doc.add_paragraph("Mobil uyumlu form elemanları", style='List Bullet')
    
    # ============ SON SAYFA ============
    doc.add_paragraph()
    doc.add_paragraph()
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("─" * 50)
    
    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer2.add_run(f"Son Güncelleme: {datetime.datetime.now().strftime('%d %B %Y')}\n")
    footer2.add_run("Dokümantasyon Sürümü: 1.0\n\n")
    footer2.add_run("Geliştirici: Muhammed Ali Aral")
    
    # Kaydet
    output_path = "KUTUPHANE_OTOMASYON_DOKUMANTASYONU.docx"
    doc.save(output_path)
    print(f"✅ Dokümantasyon başarıyla oluşturuldu: {output_path}")
    return output_path

if __name__ == "__main__":
    create_documentation()
